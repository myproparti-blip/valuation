from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_elem = OxmlElement(f'w:{edge}')
        edge_elem.set(qn('w:val'), 'single')
        edge_elem.set(qn('w:sz'), '12')
        edge_elem.set(qn('w:space'), '0')
        edge_elem.set(qn('w:color'), '000000')
        tcBorders.append(edge_elem)
    tcPr.append(tcBorders)

# Create Document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Header
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("To;")
run.bold = True
p.add_run("\nGujarat Gramin Bank , Vadodara")
run = p.add_run("\nManjalpur Branch")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(0)
p.add_run("File No: 06GGB1025 10\nDate: 31-Oct-2025")

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("VALUATION REPORT(IN RESPECT OF FLAT/HOUSE/INDUSTRIAL/SHOP)")
run.bold = True
run.font.size = Pt(11)

# GENERAL Section
general_heading = doc.add_paragraph()
run = general_heading.add_run("GENERAL")
run.bold = True
run.font.size = Pt(10)

# General Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S.No'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Details'

general_data = [
    ['1', 'Purpose for which valuation is made', 'Financial Assistance for loan from GGB Bank'],
    ['2', '(a) Date of inspection\n(b) Date on which valuation is made\nList of documents produced for pursual\n(1) Mortgage Deed :\n(2) Mortgage Deed Between :\n(3) Previous Valuation Report:\n(4) Previous Valuation Report In Favor of:\n(5) Approved Plan No:', '30-Oct-2025\n31-Oct-2025\n\nReg. No. 7204, Dated: 28/05/2025\nHemanshu Haribhai Patel & GGB Manjalpur Branch - Mr. Sanjaykumar\nIssued By I S Associates Pvt. Ltd. On Dated: 20/03/2025\nMr. Hemanshu Haribhai Patel\nApproved by Vadodara Municpal Corporation, Ward No. 4, Order No.: RAH-SHB/19/20-21, Date: 26/11/2020'],
    ['3', 'Name of the Owner/Applicant:', 'Hemanshu Haribhai Patel'],
    ['4', 'Brief description of Property', 'It is a 3bhk Residential Flat at 5th Floor of Tower A of Brookfieldz Devbhumi Residency, Flat No. A/503.'],
    ['5', 'Location of the property\n(a) Plot No/Survey No/Block No\n(b) Door/Shop No\n(c) TP Np/Village\n(d) Ward/Taluka\n(e) Mandal/District\n(f) Date of issue & Validity of layout plan\n(g) Approved map/plan issuing authority\n(h) weather genuineness or authenticity of approved map/plan verified\n(i) Any other comments by valuer on authentic of approved plan', 'R.S. No. 101, 102/2, 106/2 Paiki 2, T.P.S. No. 29, F.P. No. 9+24, At: Manjalpur, Sub District & District: Vadodara.\nFlat No. A/503, 5th Floor, Tower A, "Brookfieldz Devbhumi Residency", Nr. Tulsidham Cross Road, Manjalpur, Vadodara - 390020.\nManjalpur\nVadodara\nVadodara\n26-Nov-2020\nVadodara Municipal Corporation\nOriginal Documents Not Produced To the Valuer For Scrutinity. We have verified scan copy of original.\nProperty is constructed as per approved plan'],
    ['6', 'Postal address of the property', 'Flat No. A/503, 5th Floor, Tower A, "Brookfieldz Devbhumi Residency", Nr. Tulsidham Cross Road, Manjalpur, Vadodara - 390020.'],
    ['7', 'City/Town\nResidential Area\nComercial Area\nIndustrial Area', 'Vadodara\nYes\nNo\nNo'],
]

for row_data in general_data:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# More details (9-16)
doc.add_paragraph()

more_data = [
    ['9', 'Classification Of The Area\n(a) High/Middle/Poor\n(b) Urban/Semi Urban/Rural', 'Middle Class Area\nUrban'],
    ['10', 'Coming under Corporation limits/Village Panchayat/Municipality', 'Vadodara Municipal Corporation'],
    ['11', 'Weather convered under any State/Central Govt.enactments', 'As Per General Development Control Regulation.'],
    ['12', 'Boundaries of the property\nEast\nWest\nNorth\nSouth', 'As per Document | As per Actual\nTower B | Tower B\nStaircase, Passage | Staircase, Passage\n36 Mt. Wide Road | 36 Mt. Wide Road\nFlat No. 501, Tower B | Flat No. 501, Tower B'],
    ['13', 'Extent of the Site', 'Built Up Area (Sq.mt.): NA | Carpet Area (Sq.mt.): 68.93 | UDSL (Sq.Mt.): 20.49'],
    ['14', 'Latitude,Longitude & Co ordinates of flat', '22°16\'13.5"N 73°11\'41.8"E'],
    ['15', 'Extent of the Site Considered for valuation', 'Carpet Area (Sq.mt.): 68.93'],
    ['16', 'Weather Occupied by owner/tenant? If occupied by tenant,science how long?', 'Vacant'],
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
for row_data in more_data:
    row_cells = table2.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# II. APARTMENT BUILDING
doc.add_paragraph()
apt_heading = doc.add_paragraph()
run = apt_heading.add_run("II. APARTMENT BUILDING")
run.bold = True
run.font.size = Pt(10)

apt_data = [
    ['1', 'Nature of Apartment', 'Residential Flat'],
    ['2', 'Location\nSurvey/Block No.\nTP, FP No.\nVillage/Municipality/Corporation\nDoor No,Street or Road', 'Vadodara\nR.S. No. 101, 102/2, 106/2 Paiki 2\nT.P.S. No. 29, F.P. No. 3+24\nVadodara Municipal Corporation\n390011'],
    ['3', 'Description of the locality\nResidential/Commercial/Mixed', 'Residential Flat in Developed Area.'],
    ['4', 'Commencement Year of construction', '2025'],
    ['5', 'Number of Floor', 'Basement + Ground Floor + 7 Upper Floors'],
    ['6', 'Type Of Structure', 'RCC Structure'],
    ['7', 'Number of Dwelling units in the building', 'As Per Plan'],
    ['8', 'Quality of Construction', 'Standard'],
    ['9', 'Apperance of the building', 'Good'],
    ['10', 'Maintenance of building', 'Good'],
    ['11', 'Facilities Available\nLift\nProtected Water Supply\nUnder ground sewerage\ncar parking-Open/Covered\nis compound wall Existing?\nIs pavement laid around the building?', 'Yes\nYes\nYes\nYes\nYes\nYes'],
]

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
for row_data in apt_data:
    row_cells = table3.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# III. Flat
doc.add_paragraph()
flat_heading = doc.add_paragraph()
run = flat_heading.add_run("III. Flat")
run.bold = True
run.font.size = Pt(10)

flat_data = [
    ['1', 'The floor on which the Flat is situated', '5th Floor'],
    ['2', 'Door No, Of the Flat', 'Flat No. A-503'],
    ['3', 'Specification of the Flat\nRoof\nFlooring\nDoors\nWindows\nFittings\nFinishing', '3BHK Residential Flat\nRCC Slab\nVitrified Tiles\nWooden Framed Flush Door\nSection Windows\nGood\nInterior Finishing'],
    ['4', 'House Tax\nAssessment no\nTax paid in the name of\nTax amount', 'NA\nNA\nNA\nNA'],
    ['5', 'Electricity service connection no.\nMeter card is in name of', 'NA\nNA'],
    ['6', 'How is the maintenance of the Flat?', 'Well Maintained'],
    ['7', 'Sale Deed in the name of', 'Hemanshu Haribhai Patel'],
    ['8', 'What is the undivided area of land as per sale deed? (sq.mt.)', '20.49'],
    ['9', 'What is the plinth area of the Flat ?', 'Built Up Area (Sq.mt.): NA | Carpet Area (Sq.mt.): 68.93'],
    ['10', 'What is the FSI?', '2.7'],
    ['11', 'What is the Carpet Area of the Flat consider for valuation?', '68.93'],
    ['12', 'Is it posh/ I class/Medium / Ordinary', 'Medium'],
    ['13', 'IS It being used for residential or comercial purpose?', 'Used As Residential Flat'],
    ['14', 'is it is owner occupied or Rent out?', 'Vacant'],
    ['15', 'If rented ,what is the monthly rent?', 'Not Applicable'],
]

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
for row_data in flat_data:
    row_cells = table4.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# IV MARKETIBILITY
doc.add_paragraph()
market_heading = doc.add_paragraph()
run = market_heading.add_run("IV MARKETIBILITY")
run.bold = True
run.font.size = Pt(10)

market_data = [
    ['1', 'How is marketability?', 'Good'],
    ['2', 'What are the factors favouring for an extra potential value?', 'Prposed Fully Developed Scheme'],
    ['3', 'Any negative factors are observed which affect the market value in general?', 'The Unforeseen Events'],
]

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Light Grid Accent 1'
for row_data in market_data:
    row_cells = table5.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# V RATE
doc.add_paragraph()
rate_heading = doc.add_paragraph()
run = rate_heading.add_run("V RATE")
run.bold = True
run.font.size = Pt(10)

p1 = doc.add_paragraph('1', style='List Number')
p1.add_run('After analysing the comparable sale instances, what is the composite rate for a similar flat with same specifications in the adjoining locality?')

p1_ans = doc.add_paragraph('The estimate of Fair Market Value is based on situation, location, size, shape, road width, Neighborhood, accessibility, frontage, environmental aspects, demand and supply. The property rate is considered after information received by surrounding property holders. Also necessary information has been collected from nearby occupant. Our market inquiry among nearby occupant has revealed that similar sized property in the vicinity of the subject property is available in a range from Rs. 60000-65000 per sq. Mt. based on Carpet area.')

# VI COMPOSITE RATE
doc.add_paragraph()
composite_heading = doc.add_paragraph()
run = composite_heading.add_run("VI COMPOSITE RATE ADOPTED AFTER DEPRECIATION")
run.bold = True
run.font.size = Pt(10)

composite_data = [
    ['', 'Depreciated building rate', 'Consider In Valuation'],
    ['', 'Replacement cost of Flat with services', 'Consider In Valuation'],
    ['', 'Age of the building', '0 Years'],
    ['', 'Life of the building estimated', '50 Years'],
    ['', 'Depreciation % assuming the salvage value as 10%', 'N.A.'],
    ['', 'Depreciated ratio of the building', 'N.A.'],
    ['b', 'Total Composite rate arrived for valuation', '₹ 64,580.00 (Per Sq.mt. Carpet Area)'],
]

table_comp = doc.add_table(rows=1, cols=3)
table_comp.style = 'Light Grid Accent 1'
for row_data in composite_data:
    row_cells = table_comp.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]

# DETAILS OF VALUATION
doc.add_paragraph()
details_heading = doc.add_paragraph()
run = details_heading.add_run("DETAILS OF VALUATION")
run.bold = True
run.font.size = Pt(10)

details_table = doc.add_table(rows=1, cols=4)
details_table.style = 'Light Grid Accent 1'
hdr = details_table.rows[0].cells
hdr[0].text = 'No.'
hdr[1].text = 'DESCRIPTION'
hdr[2].text = 'Area in Sq. mt.'
hdr[3].text = 'RATE'

details = [
    ['1', 'Present value of the Flat - Carpet Area', '68.93', '₹ 64,580.00'],
    ['', 'Value Of The Flat', '', '₹ 44,51,499.40'],
    ['2', 'Fixed Furniture & Fixtures', '', '₹ 15,00,000.00'],
    ['', 'Total Value Of The Flat', '', '₹ 59,51,499.40'],
    ['', 'In Words Fifty Nine Lac Fifty One Thousand Four Hundred & Ninety Nine Rupees Only.', '', ''],
]

for row_data in details:
    row_cells = details_table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
    row_cells[3].text = row_data[3]

# Results Table
doc.add_paragraph()
results_heading = doc.add_paragraph()
run = results_heading.add_run("As a result of my appraisal and analysis,")
run.bold = True

result_table = doc.add_table(rows=1, cols=2)
result_table.style = 'Light Grid Accent 1'
result_rows = [
    ['Fair Market Market Value', '₹ 59,51,499.40'],
    ['Realizeable Value 95% of M.V', '₹ 56,63,924.43'],
    ['Distress value 80% of M.V', '₹ 47,61,199.52'],
    ['Sale Deed Value', 'NA'],
    ['Jantri Value', '₹ 16,12,962.00'],
    ['Insurable Value', '₹ 20,83,024.79'],
    ['Remarks: Rate is given on Carpet Area.', ''],
    ['Copy Of Document Shown To Us', 'Mortgage Deed, Approved Plan, Previous Valuation Report'],
]

for row_data in result_rows:
    row_cells = result_table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# STATEMENT OF LIMITING CONDITIONS
doc.add_paragraph()
limiting_heading = doc.add_paragraph()
run = limiting_heading.add_run("STATEMENT OF LIMITING CONDITIONS")
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0, 0, 255)

conditions = [
    "If this property is offered for collateral security the concerned financial institution is requested to obtained latest title report from advocate of said property.",
    "No responsibility is to be assumed for matter legal in nature nor is opinion of title rendered by this report, good title is assumed.",
    "Scope of this report is only to access present market value of the property for specific purpose, date & place. It therefore varies with purpose, period, and location, identification of rightful owner of the property, genuineness of the title deed, encumbrance if any on the property etc. be examined by the (Financial Institution) concerned authority.",
    "Possession of the any copy of this report does not carry with it the right of publication, nor any be used for any purpose by any one, except the addressee and the property owner, without the previous written consent of the appraiser, and in any event, only may be revealed in its entirety.",
    "Credibility of buyer and seller is fully responsible of financial institute. Identification of buyer & seller is from financial institute only.",
    "If found any typo error in this report is not counted for any legal action and obligation."
]

for condition in conditions:
    p = doc.add_paragraph(condition, style='List Bullet')

# VIII DECLARATION
doc.add_paragraph()
decl_heading = doc.add_paragraph()
run = decl_heading.add_run("VIII DECLARATION")
run.bold = True
run.font.size = Pt(10)

decl_table = doc.add_table(rows=1, cols=2)
decl_table.style = 'Light Grid Accent 1'

decl_data = [
    ['', 'I hereby declare that-'],
    ['a', 'I declare that I am not associated with the builder or with any of his associate companies or with the borrower directly or indirectly in the past or in the present and this report has been prepared by me with highest professional integrity.'],
    ['b', 'I further declare that I have personally inspected the site and building on 30th October, 2025.'],
    ['c', 'I further declare that all the above particulars and information given in this report are true to the best of my knowledge and belief.'],
    ['d', 'Future life of property is based on proper maintenance of the property'],
]

for row_data in decl_data:
    row_cells = decl_table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]

# Signature section
doc.add_paragraph()
sig_p = doc.add_paragraph()
sig_p.paragraph_format.left_indent = Inches(0)
sig_p.add_run("Place: Vadodara").bold = False
sig_p2 = doc.add_paragraph()
sig_p2.paragraph_format.left_indent = Inches(3.5)
sig_p2.add_run("SIGNATURE OF THE VALUER").bold = True

sig_p3 = doc.add_paragraph()
sig_p3.add_run("Date: 31/10/2025").bold = False
sig_p4 = doc.add_paragraph()
sig_p4.paragraph_format.left_indent = Inches(3.5)
sig_p4.add_run("MAHIM ARCHITECTS").bold = True

# Enclosures
doc.add_paragraph()
encl = doc.add_paragraph()
run = encl.add_run("Enclsd: 1. Declaration from the valuer")
run.bold = True

encl_text = doc.add_paragraph("The undersigned has inspected the property detailed in the Valuation report dated-30/10/2025. We are satisfied that the fair and reasonable market value of the property is Rs. 59,51,499.40/- (In Words Fifty Nine Lac Fifty One Thousand Four Hundred Ninety Nine Rupees Only).")

sig_p5 = doc.add_paragraph()
sig_p5.paragraph_format.left_indent = Inches(3.5)
sig_p5.add_run("SIGNATURE").bold = True

sig_p6 = doc.add_paragraph()
sig_p6.paragraph_format.left_indent = Inches(3.5)
sig_p6.add_run("NAME OF BRANCH OFFICIAL WITH SEAL").bold = True

# Save as docx
doc.save('Valuation_Report.docx')
print("DOCX generated successfully: Valuation_Report.docx")

# Convert DOCX to PDF using python-docx2pdf if available
try:
    from docx2pdf.convert import convert
    convert('Valuation_Report.docx', 'Valuation_Report.pdf')
    print("PDF generated successfully: Valuation_Report.pdf")
except ImportError:
    print("docx2pdf not installed. Please install with: pip install docx2pdf")
    print("Or use online converter to convert the DOCX to PDF")
